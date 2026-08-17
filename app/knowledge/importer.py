"""文档导入：读取 → 切片 → 入库 → 向量化。"""

from pathlib import Path

from ..database import db, get_connection
from .chunker import parse_markdown_units, refine_units
from .crud import create_unit, set_unit_permissions
from .embedder import get_embedder
from .vectorstore import get_vectorstore


def import_text(
    title: str,
    content: str,
    creator_id: int,
    security_level: str = "internal",
    data_domain: str = "",
    scope: list[dict] | None = None,
    source_file_name: str = "",
    file_type: str = "md",
    file_size: int = 0,
) -> dict:
    content = (content or "").strip()
    if not content:
        return {"imported": 0, "unit_ids": [], "error": "内容为空"}

    if content.startswith("# "):
        units = refine_units(parse_markdown_units(content))
    else:
        units = refine_units([{"title": title, "content": content}])

    # R6 幂等去重：文件名 + 首个单元内容一致视为重复导入（内容哈希等价）
    if source_file_name and units:
        with get_connection() as conn:
            dup = conn.execute(
                "SELECT COUNT(*) c FROM knowledge_units WHERE source_file_name=? AND content=?",
                (source_file_name, units[0]["content"]),
            ).fetchone()["c"]
        if dup:
            return {
                "imported": 0,
                "unit_ids": [],
                "duplicated": True,
                "error": "重复导入（文件名与内容一致）",
            }

    embedder = get_embedder()
    store = get_vectorstore()
    texts = [u["content"] for u in units]
    vectors = embedder.embed(texts)
    ids: list[str] = []
    titles = []
    with db() as conn:
        for u in units:
            uid = create_unit(
                conn,
                title=u["title"],
                content=u["content"],
                category=data_domain or "默认",
                source_file_name=source_file_name,
                file_type=file_type,
                file_size=file_size,
                security_level=security_level,
                data_domain=data_domain,
                creator_id=creator_id,
            )
            set_unit_permissions(
                conn, uid, scope or [{"target_type": "user", "target_id": str(creator_id)}]
            )
            ids.append(str(uid))
            titles.append(u["title"])
    store.add(ids, texts, vectors, [{"unit_id": i, "title": t} for i, t in zip(ids, titles)])
    return {"imported": len(units), "unit_ids": [int(i) for i in ids]}


def import_file(
    path,
    creator_id: int,
    security_level: str = "internal",
    data_domain: str = "",
    scope: list[dict] | None = None,
    source_name: str | None = None,
) -> dict:
    p = Path(path)
    ext = p.suffix.lower()
    text = ""
    error = ""
    try:
        if ext in (".md", ".txt"):
            text = p.read_text(encoding="utf-8", errors="ignore")
        elif ext == ".pdf":
            try:
                import pdfplumber

                with pdfplumber.open(p) as pdf:
                    text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            except Exception as e:  # noqa: BLE001
                error = f"PDF 解析失败：{e}"
        elif ext == ".docx":
            try:
                from docx import Document

                doc = Document(str(p))
                parts: list[str] = []
                for para in doc.paragraphs:
                    txt = para.text.strip()
                    if not txt:
                        continue
                    # 将 Word 标题样式转为 Markdown 标题，便于后续按标题拆分为知识单元
                    style = (para.style.name or "") if para.style else ""
                    if style.lower().startswith("heading "):
                        try:
                            level = int(style.split(" ", 1)[1])
                        except ValueError:
                            level = 1
                        parts.append("#" * max(1, min(level, 6)) + " " + txt)
                    elif style.lower() == "title":
                        parts.append("# " + txt)
                    else:
                        parts.append(txt)
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        line = " | ".join(c for c in cells if c)
                        if line:
                            parts.append(line)
                text = "\n".join(parts)
            except Exception as e:  # noqa: BLE001
                error = f"Word 解析失败：{e}"
        else:
            error = f"不支持的文件类型：{ext}"
    except Exception as e:  # noqa: BLE001
        error = f"读取失败：{e}"

    if error or not text.strip():
        return {"imported": 0, "unit_ids": [], "error": error or "文件为空"}
    display = source_name or p.name  # 落盘名为唯一 uuid 前缀时，展示/去重用原始文件名
    return import_text(
        title=Path(display).stem,
        content=text,
        creator_id=creator_id,
        security_level=security_level,
        data_domain=data_domain,
        scope=scope,
        source_file_name=display,
        file_type=Path(display).suffix.lstrip("."),
        file_size=p.stat().st_size,
    )
